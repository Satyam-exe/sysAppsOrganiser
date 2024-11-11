import os

import subprocess

from docx import Document

import functions


TEMPLATE_PATH = './template/template.docx'
FINAL_PATH = './data/final/'


def create_docx(email_id):
    doc = Document(TEMPLATE_PATH)

    data = {
        'name_of_applicant': functions.name_from_id(email_id),
        'class_section_of_applicant': functions.class_section_from_id(email_id),
        'post_of_applicant': functions.post_from_id(email_id),
    }

    question_answers = functions.responses_from_id(email_id)

    for paragraph in doc.paragraphs:
        for placeholder, value in data.items():
            if f"{{{{ {placeholder} }}}}" in paragraph.text:
                for run in paragraph.runs:
                    if f"{{{{ {placeholder} }}}}" in run.text:
                        run.text = run.text.replace(f"{{{{ {placeholder} }}}}", value)

        if '{{ question_answers }}' in paragraph.text:
            for run in paragraph.runs:
                if '{{ question_answers }}' in run.text:
                    run.text = run.text.replace('{{ question_answers }}', '')

            for question, answer in question_answers.items():
                paragraph.add_run(f"{question}: \n").bold = True
                paragraph.add_run(f"{answer}\n\n")

    doc.save(f"{FINAL_PATH}/{functions.name_from_id(email_id)}.docx")


def convert_docx_to_pdf(email_id):
    docx_path = f"{FINAL_PATH}/{functions.name_from_id(email_id)}.docx"
    pdf_path = f"{FINAL_PATH}/{functions.name_from_id(email_id)}.pdf"

    subprocess.run(
        [r"C:\Program Files\LibreOffice\program\soffice.exe", "--headless", "--convert-to", "pdf", "--outdir",
         os.path.dirname(pdf_path), docx_path])

    if os.path.exists(pdf_path):
        os.remove(docx_path)